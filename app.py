import streamlit as st
import numpy as np
import io
from reportlab.pdfgen import canvas
from docx import Document

# --- FUNÇÕES CORE DE ENGENHARIA (NBR 17227:2025) ---

def calc_ia_step(ibf, g, k):
    """Equação 1: Corrente de Arco Intermediária"""
    k1, k2, k3, k4, k5, k6, k7, k8, k9, k10 = k
    log_base = k1 + k2 * np.log10(ibf) + k3 * np.log10(g)
    poli = (k4*ibf**6 + k5*ibf**5 + k6*ibf**4 + k7*ibf**3 + k8*ibf**2 + k9*ibf + k10)
    return 10**(log_base * poli)

def calc_var_cf(v_oc):
    """Equação 2: Fator de Variação da Corrente"""
    k = [0, 0, 0, 0, -0.0001, 0.0022, 0.02] # Tabela 5
    return (k[0]*v_oc**6 + k[1]*v_oc**5 + k[2]*v_oc**4 + k[3]*v_oc**3 + k[4]*v_oc**2 + k[5]*v_oc + k[6])

def calc_cf(dimensoes, tipo_inv):
    """Equações 13-15: Fator de Correção do Invólucro"""
    A, L, P = [d/25.4 for d in dimensoes] # Conversão para pol
    ees = (A + L) / 2.0 if P > 8 else A
    b1, b2, b3 = -0.0003, 0.03441, 0.4325 # VCB Típico
    cf = b1*ees**2 + b2*ees + b3
    return (cf if tipo_inv == "Típico" else 1.0/cf), ees

def calc_en_step(ia, ibf, g, d, t, k, cf):
    """Equação 3: Energia Incidente Intermediária (J/cm²)"""
    k1, k2, k3, k4, k5, k6, k7, k8, k9, k10, k11, k12, k13 = k
    poli_den = (k4*ibf**7 + k5*ibf**6 + k6*ibf**5 + k7*ibf**4 + k8*ibf**3 + k9*ibf**2 + k10*ibf)
    termo_ia = (k3 * ia) / poli_den if poli_den != 0 else 0
    exp = (k1 + k2*np.log10(g) + termo_ia + k11*np.log10(ibf) + k12*np.log10(d) + k13*np.log10(ia) + np.log10(1.0/cf))
    return (12.552 / 50.0) * t * (10**exp)

def calc_dla_step(ia, ibf, g, t, k, cf):
    """Equações 7-10: Distância Limite de Arco Intermediária (mm)"""
    k1, k2, k3, k4, k5, k6, k7, k8, k9, k10, k11, k12, k13 = k
    poli_den = (k4*ibf**7 + k5*ibf**6 + k6*ibf**5 + k7*ibf**4 + k8*ibf**3 + k9*ibf**2 + k10*ibf)
    termo_ia = (k3 * ia) / poli_den if poli_den != 0 else 0
    log_fixo = (k1 + k2*np.log10(g) + termo_ia + k11*np.log10(ibf) + k13*np.log10(ia) + np.log10(1.0/cf))
    log_d = (np.log10(5.0 / ((12.552 / 50.0) * t)) - log_fixo) / k12
    return 10**log_d

def interpolar(v, v1, v2, v3, f1, f2, f3):
    """Equações 16-33: Interpolação Final de Tensão"""
    if v <= 0.6: return f1
    if v <= 2.7: return f1 + (f2 - f1) * (v - 0.6) / 2.1
    return f2 + (f3 - f2) * (v - 2.7) / 11.6

# --- INTERFACE E FLUXO ---

def main():
    st.set_page_config(page_title="NBR 17227 Expert", layout="wide")
    st.title("⚡ Gestão de Risco de Arco Elétrico - NBR 17227:2025")

    # Bancos de Dados Normativos
    equipamentos = {
        "CCM 15 kV": [152, 914.4, 914.4, 914.4, 914.4],
        "Conjunto de manobra 15 kV": [152, 914.4, 1143.0, 762.0, 762.0],
        "CCM 5 kV": [104, 914.4, 660.4, 660.4, 660.4],
        "CCM e painel típico de BT": [25, 457.2, 355.6, 304.8, 203.3]
    }
    
    tab1, tab2, tab3 = st.tabs(["📏 Tabela 1 (Equipamentos)", "🧪 Cálculos NBR", "📄 Relatório"])

    with tab1:
        st.header("Consulta de Dimensões Normativas")
        escolha = st.selectbox("Selecione o Equipamento:", list(equipamentos.keys()))
        info = equipamentos[escolha]
        c = st.columns(5)
        for i, t in enumerate(["GAP", "D_trab", "Alt (A)", "Larg (L)", "Prof (P)"]):
            c[i].metric(t, f"{info[i]} mm")

    with tab2:
        with st.form("calc_form"):
            st.subheader("Parâmetros de Entrada")
            c1, c2, c3 = st.columns(3)
            with c1:
                v_oc = st.number_input("Tensão Voc (kV)", value=13.80)
                i_bf = st.number_input("Curto-Circuito Ibf (kA)", value=4.852)
            with c2:
                gap = st.number_input("Gap G (mm)", value=float(info[0]))
                dist = st.number_input("Distância D (mm)", value=float(info[1]))
            with c3:
                tempo = st.number_input("Duração T (ms)", value=488.0)
                tipo_inv = st.radio("Invólucro:", ["Típico", "Raso"])
            
            calc_btn = st.form_submit_button("Calcular Resultados Finais")

        if calc_btn:
            # Coeficientes Tabela 4 (Ia) e 6 (Energia/DLA) - Configuração VCB
            k_ia = {600: [-0.04287, 1.035, -0.083, 0, 0, -4.783e-9, 1.962e-6, -0.000229, 0.003141, 1.092],
                    2700: [0.0065, 1.001, -0.024, -1.557e-12, 4.556e-10, -4.186e-8, 8.346e-7, 5.482e-5, -0.003191, 0.9729],
                    14300: [0.005795, 1.015, -0.011, -1.557e-12, 4.556e-10, -4.186e-8, 8.346e-7, 5.482e-5, -0.003191, 0.9729]}
            
            k_en = {600: [0.753364, 0.566, 1.752636, 0, 0, -4.783e-9, 1.962e-6, -0.000229, 0.003141, 1.092, 0, -1.598, 0.957],
                    2700: [2.40021, 0.165, 0.354202, -1.557e-12, 4.556e-10, -4.186e-8, 8.346e-7, 5.482e-5, -0.003191, 0.9729, 0, -1.569, 0.9778],
                    14300: [3.825917, 0.11, -0.999749, -1.557e-12, 4.556e-10, -4.186e-8, 8.346e-7, 5.482e-5, -0.003191, 0.9729, 0, -1.568, 0.99]}

            # Processamento
            cf, ees = calc_cf(info[2:], tipo_inv)
            ia600, ia2700, ia14300 = [calc_ia_step(i_bf, gap, k_ia[v]) for v in [600, 2700, 14300]]
            e600, e2700, e14300 = [calc_en_step(ia, i_bf, gap, dist, tempo, k_en[v], cf) for ia, v in zip([ia600, ia2700, ia14300], [600, 2700, 14300])]
            dla600, dla2700, dla14300 = [calc_dla_step(ia, i_bf, gap, tempo, k_en[v], cf) for ia, v in zip([ia600, ia2700, ia14300], [600, 2700, 14300])]
            
            # Interpolação Final
            ia_f = interpolar(v_oc, ia600, ia2700, ia14300)
            e_f_cal = interpolar(v_oc, e600, e2700, e14300) / 4.184
            dla_f = interpolar(v_oc, dla600, dla2700, dla14300)
            var_cf = calc_var_cf(v_oc)
            ia_min = ia_f * (1 - 0.5 * var_cf)

            # Vestimenta
            cat = "Cat 2 (8 cal)" if e_f_cal <= 8 else ("Cat 4 (40 cal)" if e_f_cal <= 40 else "PERIGO")
            
            st.session_state['res'] = {"Iarc": ia_f, "E": e_f_cal, "DLA": dla_f, "Imin": ia_min, "Cat": cat, "Voc": v_oc}
            
            st.divider()
            r1, r2, r3, r4 = st.columns(4)
            r1.metric("I_arc Final (kA)", f"{ia_f:.5f}")
            r2.metric("Energia (cal/cm²)", f"{e_f_cal:.5f}")
            r3.metric("DLA Final (mm)", f"{dla_f:.1f}")
            r4.metric("I_arc Mínima (kA)", f"{ia_min:.5f}")
            st.warning(f"**Vestimenta recomendada:** {cat}")

    with tab3:
        if 'res' in st.session_state:
            st.write("Selecione o formato para exportação do laudo técnico:")
            
            # Gerador PDF
            def export_pdf():
                buf = io.BytesIO()
                c = canvas.Canvas(buf)
                c.drawString(100, 800, f"Laudo de Arco Elétrico - NBR 17227")
                c.drawString(100, 780, f"Tensão: {st.session_state['res']['Voc']} kV")
                c.drawString(100, 760, f"Energia Incidente: {st.session_state['res']['E']:.4f} cal/cm2")
                c.drawString(100, 740, f"DLA: {st.session_state['res']['DLA']:.1f} mm")
                c.drawString(100, 720, f"Vestimenta: {st.session_state['res']['Cat']}")
                c.save()
                return buf.getvalue()

            # Gerador Word
            def export_word():
                doc = Document()
                doc.add_heading('Laudo Técnico de Arco Elétrico', 0)
                doc.add_paragraph(f"Tensão do Sistema: {st.session_state['res']['Voc']} kV")
                doc.add_paragraph(f"Energia Calculada: {st.session_state['res']['E']:.4f} cal/cm2")
                doc.add_paragraph(f"Vestimenta Obrigatória: {st.session_state['res']['Cat']}")
                buf = io.BytesIO()
                doc.save(buf)
                return buf.getvalue()

            st.download_button("Baixar Laudo (PDF)", export_pdf(), "laudo.pdf", "application/pdf")
            st.download_button("Baixar Laudo (Word)", export_word(), "laudo.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.info("Realize um cálculo na aba anterior para gerar o relatório.")

if __name__ == "__main__":
    main()
